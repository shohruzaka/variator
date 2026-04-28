"""Word (docx) fayllarga eksport qilish (7-bosqich).

Bu modul `Variant` obyektlarini o'qib, ularni foydalanuvchiga taqdim
etishga tayyor `.docx` fayllarga aylantiradi va `output/` papkasiga saqlaydi.
Faqat fayl yozish bilan shug'ullanadi.
"""

import os
from pathlib import Path
from typing import Callable

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.models import Variant


def _setup_document_format(doc: Document, font_size: int = 12):
    """Hujjatni albom ko'rinishiga, 3 kalonkaga va belgilangan shriftga o'tkazadi."""
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(font_size)
    style.paragraph_format.space_after = Pt(0)  # Paragraflar orasidagi standart bo'shliqni olib tashlash
    style.paragraph_format.line_spacing = 1.0   # Qatorlar orasini (line spacing) 1.0 qilib zichlash
    
    for section in doc.sections:
        # Hujjatni Albom (Landscape) ko'rinishiga o'tkazish
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        
        # Hoshiyalarni (margin) qisqartirish
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # 3 ta ustunga (kalonka) ajratish
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')
        if not cols:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        else:
            cols = cols[0]
            
        cols.set(qn('w:num'), '3')
        cols.set(qn('w:space'), '720')


def _add_answer_table(doc: Document, num_questions: int):
    """Javoblarni belgilash uchun jadval qo'shadi."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Javoblar:")
    run.bold = True
    run.font.name = 'Times New Roman'
    
    chunk_size = 10  # Kalonkaga (ustunga) chiroyli sig'ishi uchun 10 tadan bo'lamiz
    for i in range(0, num_questions, chunk_size):
        chunk = min(chunk_size, num_questions - i)
        table = doc.add_table(rows=2, cols=chunk)
        table.style = 'Table Grid'
        
        table.rows[0].height = Pt(18)
        table.rows[1].height = Pt(24) # O'quvchi javob yozadigan qatorni kattaroq qildik
        
        for j in range(chunk):
            q_num = i + j + 1
            cell_top = table.cell(0, j)
            cell_top.text = str(q_num)
            cell_top.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell_top.paragraphs[0].runs[0].font.size = Pt(10)
            cell_top.paragraphs[0].runs[0].bold = True
            cell_top.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell_bottom = table.cell(1, j)
            cell_bottom.text = " "
            cell_bottom.paragraphs[0].runs[0].font.size = Pt(10)
            cell_bottom.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(2)

def _add_formatted_runs(paragraph, text: str):
    """Matn ichidagi `kod` qismlarini ajratib, alohida shrift (Consolas) beradi."""
    parts = text.split('`')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # ` ichidagi qism (toq indekslar)
            run.font.name = 'Consolas'


def _add_document_header(
    doc: Document,
    subject_name: str,
    assessment_type: str,
    variant_label: str,
) -> None:
    """Hujjat tepasiga sarlavha va talaba ma'lumotlarini chiroyli formatda yozadi."""
    
    # 1. Variant sarlavhasi (O'rtada, Katta o'lchamda)
    p_variant = doc.add_paragraph()
    p_variant.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_variant.paragraph_format.space_before = Pt(0)
    p_variant.paragraph_format.space_after = Pt(8)
    run_variant = p_variant.add_run(variant_label)
    run_variant.bold = True
    run_variant.font.name = 'Times New Roman'
    run_variant.font.size = Pt(14)

    # 2. Fan va Nazorat turi
    if subject_name:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_label = p.add_run("Fan: ")
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_value = p.add_run(subject_name)
        run_value.font.name = 'Times New Roman'

    if assessment_type:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_label = p.add_run("Nazorat turi: ")
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_value = p.add_run(assessment_type)
        run_value.font.name = 'Times New Roman'

    # 3. Talaba ma'lumotlari (F.I.Sh, Guruh, Sana)
    p_student = doc.add_paragraph()
    p_student.paragraph_format.space_before = Pt(6)
    p_student.paragraph_format.space_after = Pt(4)
    run_s_label = p_student.add_run("F.I.Sh: ")
    run_s_label.bold = True
    run_s_label.font.name = 'Times New Roman'
    p_student.add_run("__________________________")

    p_info = doc.add_paragraph()
    p_info.paragraph_format.space_after = Pt(12)
    run_g_label = p_info.add_run("Guruh: ")
    run_g_label.bold = True
    run_g_label.font.name = 'Times New Roman'
    p_info.add_run("_______  ")
    
    run_d_label = p_info.add_run("Sana: ")
    run_d_label.bold = True
    run_d_label.font.name = 'Times New Roman'
    p_info.add_run("________")

def export_variants_to_docx(
    variants: list[Variant],
    output_dir: str | Path,
    font_size: int = 12,
    subject_name: str = "",
    assessment_type: str = "",
    progress_cb: Callable[[int, int], None] | None = None
) -> Path:
    """Barcha variantlarni bitta Word fayliga yozadi (har biri yangi varaqdan boshlanadi).

    Args:
        variants: Generatsiya qilingan test variantlari ro'yxati.
        output_dir: Fayl saqlanadigan papka manzili.
        font_size: Word hujjatining shrift o'lchami (standart 12).
        subject_name: Fan nomi (bo'sh bo'lsa yozilmaydi).
        assessment_type: Nazorat turi (bo'sh bo'lsa yozilmaydi).
        progress_cb: Jarayonni foizda ko'rsatish uchun callback.

    Returns:
        Yaratilgan faylning to'liq manzili.
    """
    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Formatni qo'llash (Albom, 3 ta kalonka, belgilangan shrift)
    _setup_document_format(doc, font_size)

    for idx, variant in enumerate(variants):
        _add_document_header(doc, subject_name, assessment_type, f"{variant.number}-variant")
        
        # Jadval qo'shish
        _add_answer_table(doc, len(variant.questions))
        
        for q in variant.questions:
            p_q = doc.add_paragraph()
            _add_formatted_runs(p_q, f"{q.number}. {q.text}")
            p_q.paragraph_format.space_after = Pt(2)
            for i, opt in enumerate(q.options):
                p_opt = doc.add_paragraph()
                _add_formatted_runs(p_opt, f"{opt.letter}) {opt.text}")
                p_opt.paragraph_format.space_after = Pt(8) if i == len(q.options) - 1 else Pt(0)
                
        if idx < len(variants) - 1:
            doc.add_page_break()  # Keyingi variantni yangi sahifadan boshlash
            
        if progress_cb:
            progress_cb(idx + 1, len(variants))
            
    file_path = out_path / "Barcha_variantlar.docx"
    doc.save(str(file_path))
    return file_path


def export_answers_to_docx(
    variants: list[Variant],
    output_path: str | Path,
    subject_name: str = "",
    assessment_type: str = "",
) -> Path:
    """Barcha variantlarning javoblarini bitta Word fayliga yozadi.

    Args:
        variants: Generatsiya qilingan test variantlari ro'yxati.
        output_path: Saqlanadigan Word faylining to'liq manzili.
        subject_name: Fan nomi (bo'sh bo'lsa yozilmaydi).
        assessment_type: Nazorat turi (bo'sh bo'lsa yozilmaydi).

    Returns:
        Yaratilgan faylning manzili.
    """
    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    title = doc.add_heading("Javoblar kaliti", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subject_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run_label = p.add_run("Fan: ")
        run_label.bold = True
        p.add_run(subject_name)

    if assessment_type:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run_label = p.add_run("Nazorat turi: ")
        run_label.bold = True
        p.add_run(assessment_type)

    for variant in variants:
        h2 = doc.add_heading(f"{variant.number}-variant", level=2)
        h2.paragraph_format.space_before = Pt(16)
        h2.paragraph_format.space_after = Pt(8)
        
        chunk_size = 10
        ans_len = len(variant.answer_key)
        
        for i in range(0, ans_len, chunk_size):
            chunk = min(chunk_size, ans_len - i)
            table = doc.add_table(rows=2, cols=chunk)
            table.style = 'Table Grid'
            
            table.rows[0].height = Pt(18)
            table.rows[1].height = Pt(18)
            
            for j in range(chunk):
                q_num = i + j + 1
                ans = variant.answer_key[i + j]
                
                cell_top = table.cell(0, j)
                cell_top.text = str(q_num)
                cell_top.paragraphs[0].runs[0].bold = True
                cell_top.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                cell_bottom = table.cell(1, j)
                cell_bottom.text = ans
                cell_bottom.paragraphs[0].runs[0].bold = True
                cell_bottom.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 112, 192) # Ko'k rangda
                cell_bottom.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_after = Pt(2)

    doc.save(str(out_path))
    return out_path
