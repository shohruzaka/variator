"""Exporter (Word) uchun unit testlar (7-bosqich)."""

from pathlib import Path

from docx import Document

from src.exporter_docx import export_variants_to_docx
from src.models import Option, Question, Variant


def _make_sample_variants() -> list[Variant]:
    """Test uchun 2 ta variant yaratadi."""
    q1 = Question(1, "Savol 1", [Option("A", "a1"), Option("B", "b1"), Option("C", "c1", is_correct=True), Option("D", "d1")])
    q2 = Question(2, "Savol 2", [Option("A", "a2"), Option("B", "b2"), Option("C", "c2"), Option("D", "d2", is_correct=True)])
    
    v1 = Variant(number=1, seed=101, questions=[q1, q2])
    
    q3 = Question(1, "Savol 2", [Option("A", "a2"), Option("B", "b2", is_correct=True), Option("C", "c2"), Option("D", "d2")])
    q4 = Question(2, "Savol 1", [Option("A", "a1", is_correct=True), Option("B", "b1"), Option("C", "c1"), Option("D", "d1")])
    
    v2 = Variant(number=2, seed=102, questions=[q3, q4])
    
    return [v1, v2]


def test_export_variants_creates_single_file(tmp_path: Path):
    """Eksport funksiyasi barcha variantlarni bitta `Barcha_variantlar.docx` fayliga yozishi kerak."""
    variants = _make_sample_variants()
    output_dir = tmp_path / "output"

    result_path = export_variants_to_docx(variants, output_dir)

    assert output_dir.exists()
    assert result_path.exists()
    assert result_path.name == "Barcha_variantlar.docx"


def test_export_variants_content(tmp_path: Path):
    """Yaratilgan faylda har ikkala variant sarlavhasi va savollari mavjud bo'lishi kerak."""
    variants = _make_sample_variants()
    output_dir = tmp_path / "output"

    result_path = export_variants_to_docx(variants, output_dir)

    doc = Document(str(result_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Ikkala variant ham bitta hujjatda bo'lishi kerak
    assert any("1-variant" in p for p in paragraphs)
    assert any("2-variant" in p for p in paragraphs)
    assert any("1. Savol 1" in p for p in paragraphs)
    assert any("A) a1" in p for p in paragraphs)
    assert any("D) d1" in p for p in paragraphs)
    assert any("2. Savol 2" in p for p in paragraphs)


def test_export_empty_variants_list(tmp_path: Path):
    """Bo'sh ro'yxat berilganda ham fayl yaratilishi va papka mavjud bo'lishi kerak."""
    output_dir = tmp_path / "output"
    result_path = export_variants_to_docx([], output_dir)

    assert output_dir.exists()
    assert result_path.exists()


def test_export_answers_to_docx(tmp_path: Path):
    """Javoblar kaliti Word faylga to'g'ri yozilishini tekshirish."""
    from src.exporter_docx import export_answers_to_docx
    
    variants = _make_sample_variants()
    out_file = tmp_path / "javoblar.docx"
    
    result_path = export_answers_to_docx(variants, out_file)
    
    assert result_path.exists()
    
    doc = Document(str(result_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Kutilgan format:
    # Javoblar kaliti
    # 1-variant
    # (Jadval formatida)
    
    assert "Javoblar kaliti" in paragraphs[0]
    assert "1-variant" in paragraphs[1]
    assert "2-variant" in paragraphs[2]
    
    tables = doc.tables
    assert len(tables) >= 2
    
    # 1-variant jadvali
    assert [c.text for c in tables[0].rows[0].cells] == ["1", "2"]
    assert [c.text for c in tables[0].rows[1].cells] == ["C", "D"]
    
    # 2-variant jadvali
    assert [c.text for c in tables[1].rows[0].cells] == ["1", "2"]
    assert [c.text for c in tables[1].rows[1].cells] == ["B", "A"]
